import os
import re
from config.config import readsample1
from config.config import readsample2
from config.config import readsample3
from config.config import readpathdir
from docx import Document
import config
def zyout(path,text):
    with open(path, 'w',encoding='utf-8') as f:
        print(text,file=f)
# -----------------------例文打印-------------------------
def sampleshow():
    file_path=readsample3
    doc = Document(file_path)
    paragraphs=doc.paragraphs
    print(file_path)
    print(paragraphs)


def catalogue_get(doc):
    docx = Document(doc).add_heading
    lastest_heading = 0
    record = ['1']  # 记录目录结构
    point = '.'
    for paragraph in docx.paragraphs:

        if paragraph.style.name[:7] == 'Heading':
            this_heading = int(paragraph.style.name[-1])

            result = ''.join(record) + point
            if this_heading == 1 and lastest_heading == 0:
                heading = ''.join(record) + '.'
            else:
                if this_heading > lastest_heading:
                    record.append('1')
                elif this_heading == lastest_heading:
                    record[-1] = str(int(record[-1]) + 1)
                else:
                    record[this_heading - 1] = str(int(record[this_heading - 1]) + 1)
                    record[this_heading:] = []
            heading = '.'.join(record) + point  # 显示一段目录

            print(heading, paragraph.text, paragraph.style.name, sep='   ')
            lastest_heading = this_heading
# catalogue_get(doc)
#目录提取
def menuout(path,text):
    with open(path, 'w',encoding='utf-8') as f:
        print(text,file=f)
def menuextract(paragraphs,a):
    menu_console_out=0
    if menu_console_out == 1:
        print("in ",a.getname()," menuextract")
    i=0
    menutext=[]
    for paragraph in paragraphs:
        if 'Heading ' in paragraph.style.name and paragraph.text.replace(" ", "") != '':
            if menu_console_out==1:
                print(paragraph.text, paragraph.style.name,"pi = ",i)
            menutext.append([paragraph.text, paragraph.style.name,i])
        i += 1
    menuout('D:\PyProject\docx_input\menu_out_files\\'+a.getname()+'.txt',menutext)
    if menu_console_out == 1:
        print(a.getname(), "menuextract end")
    return menutext
    # print(list(filter(None,re.split('\t| ',p.text))))

def zykeywordextract(paragraphs,article):
    name=article.getname()
    zykey_console_out=0
    if zykey_console_out == 1:
        print(article.getname() + '摘要和关键词打印开始-----------------')
    pi=0#段落index
    fzyi=0#目标index
    for i in paragraphs:
        xi=list(filter(None,i.text.split(" ")))#去除段落中的空格
        pi+=1#pi指向下一个段落
        if len(xi)>=1 and xi[0]=='摘要':

            # print(xi[0])
            fzyi=pi
        elif len(xi)>=2 and xi[0]=='摘' and xi[1]=='要':

            # print(xi[0])
            # print(xi[1])
            fzyi = pi
        if len(xi)>=1and len(xi[0]) >= 3 and xi[0][0:3] == '关键词':
            # print(xi[0])
            break
    zytext=[]
    keywordtext=[]
    for pit in paragraphs[fzyi:pi-1]:
            #每个pit.text为摘要中的一段话
        if zykey_console_out == 1:
            print(pit.text)
        zytext.append(pit.text)
    try :
        keywordstr=paragraphs[pi-1].text.split('：')[1]
    except:
        print("关键词部分不由中文冒号分隔")

    if '，' in keywordstr:
        keywordtext = list(keywordstr.split('，'))
    elif '、' in keywordstr:
        keywordtext = list(keywordstr.split("、"))
    else:
        print("------------------------")
        print("关键词不由中文 , or 、分隔")
        print("------------------------")
        # except:
        #     print("关键词不由 , or 、分隔")

    zyout('D:\PyProject\docx_input\zy_output_files\\'+name+'.txt',zytext)
    zyout('D:\PyProject\docx_input\keyword_output_files\\'+name+'.txt',keywordtext)
    if zykey_console_out == 1:
        print(article.getname()+'摘要和关键词打印结束-----------------')
    return zytext,keywordtext
#---------------------摘要和关键词打印结束-----------------------------------------------


# files=os.listdir(config.config.readpathdir)
# for file in files:
#     print(file)
#     name=file.split('.')[0]
#     file_path = os.path.join(config.config.readpathdir, file)
#     doc = Document(file_path)
#     menutext=[]
#     i=0
#     for paragraph in doc.paragraphs:
#         if 'Heading ' in paragraph.style.name and paragraph.text.replace(" ", "") != '':
#             print(paragraph.text, paragraph.style.name,"pi = ",i)
#             menutext.append([paragraph.text, paragraph.style.name,i])
#         i += 1
#     menuout('D:\PyProject\docx_input\menu_out_files\\'+name+'.txt',menutext)
#     for index  in range(0,len(menutext)):
#         if index!=len(menutext)-1:
#             for pindex in range(menutext[index][2],menutext[index+1][2]):
#                 print(doc.paragraphs[pindex].text)
#
#     break
#以下为备用目录抽取方法
    # pi=0
    # fzyi=0
    # for i in doc.paragraphs:
    #     xi=list(filter(None,re.split('\t| ',i.text)))
    #     pi+=1
    #     if len(xi)>=1 and xi[0]=='目录':
    #         print(file,'目录定位',pi)
    #         print(xi[0])
    #         fzyi=pi
    #     elif len(xi)>=2 and xi[0]=='目' and xi[1]=='录':
    #         print(file,'目   录 定位',pi)
    #         print(xi[0])
    #         print(xi[1])
    #         fzyi = pi
    #     if len(xi)>=1 and xi[0]=='绪论' or len(xi)>=2 and xi[1]=='绪论':
    #         print(file,"绪论 定位",pi)
    #         print(xi[0])
    #         fzyi=pi
    #     elif len(xi)>=2 and xi[0]=='绪' and xi[1]=='论':
    #         print(file,'绪   论 定位',pi)
    #         print(xi[0])
    #         print(xi[1])
    #         fzyi = pi
    #     if len(xi)>=1and len(xi[0]) >= 4 and xi[0][0:4] == '参考文献':
    #         print(file, '参考文献 定位',pi)
    #         print(xi[0])
    #         break


