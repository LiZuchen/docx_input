import os

import config.config
from docx import Document

def cut(str1, str2, text, out, rws):
    if text == str1:
        rws = 1
    if text == str2:
        rws = 0
    if rws == 2:
        out.append(text.replace(" ",""))
        # print(text)
    elif rws == 1:
        rws = 2
    return rws
def rws4clean(rws4out):
    rws4c=[]
    for i in rws4out:
        if i[0:5]=='计算机学院' or i[0:5]=='高等理工学'  :
            return rws4c
        else:
            rws4c.append(i)
    return rws4c
def rwsout(path,outlist):
    with open(path, 'w') as f:
        print(outlist,file=f)
def rws(paragraphs,name):
    i = 0
    rws1 = 0
    rws1out = []
    rws2 = 0
    rws2out = []
    rws3 = 0
    rws3out = []
    rws4 = 0
    rws4out = []
    for paragraph in paragraphs:
        # print(i)
        # i += 1
        rws1 = cut("Ⅰ、毕业设计（论文）题目：",
                   "Ⅱ、毕业设计（论文）使用的原始资料（数据）及设计技术要求：",
                   paragraph.text, rws1out, rws1)
        rws2 = cut("Ⅱ、毕业设计（论文）使用的原始资料（数据）及设计技术要求：",
                   "Ⅲ、毕业设计（论文）工作内容：",
                   paragraph.text, rws2out, rws2)
        rws3 = cut("Ⅲ、毕业设计（论文）工作内容：",
                   "Ⅳ、主要参考资料：",
                   paragraph.text, rws3out, rws3)
        rws4 = cut("Ⅳ、主要参考资料：",
                   "我声明，本论文及其研究工作是由本人在导师指导下独立完成的，"
                   "在完成论文时所利用的一切资料均已在参考文献中列出。",
                   paragraph.text, rws4out, rws4)
    rwslist=[]
    rwslist.append(list(filter(None,rws1out)))#题目
    rwslist.append(list(filter(None,rws2out)))#原始资料（数据）及设计技术
    rwslist.append(list(filter(None,rws3out)))#工作内容
    rws4out=rws4clean(list(filter(None,rws4out)))
    rwslist.append(list(filter(None,rws4out)))#主要参考资料

    rws_console_out=0
    if rws_console_out:
        print("-------------------------------")
        for rwsi in rwslist:
            print(rwsi)
        print("-------------------------------")
    rwsout("D:\PyProject\docx_input\\rws_output_files\\"+name+'.txt',rwslist)
    return rwslist

def rwsbuild():
    files=os.listdir(config.config.readpathdir)
    for file in files:

        print(file)
        file_path=os.path.join(config.config.readpathdir,file)
        doc = Document(file_path)
        # print(file_path)
        # print(len(doc.paragraphs))
    # 输出的是列表，列表中一共有421份内容
    # [<docx.text.paragraph.Paragraph object at 0x000001906641C100>,...
    # <docx.text.paragrap h.Paragraph object at 0x000001906643C940>]
    # 输出每一段的文字
        name=file.split('.')[0]
        rws(doc.paragraphs,name)

# paragraph1 = doc.paragraphs[1]
# runs1 = paragraph1.runs
# print(runs1)
#
# for run1 in runs1:
# 	print(run1.text)

# 输出结果：该段落为空，没有文字块
# []
