# -*- coding: utf-8 -*-
import os
import time
from article_related.article import article
from config.config import readpathdir, readpathdirall
from docx import Document
from input.block_divide import blockdivide
from input.readsentences import zykeywordextract, menuextract
from input.rwsinput import rws
import win32com.client as wc
import numpy as np

def article_build():

    batch_num=6
    readpathdiralli=readpathdirall+'\\'+str(batch_num)
    files=os.listdir(readpathdiralli)
    articlelist=[]
    for file in files:
        print(file)
        file_path = os.path.join(readpathdiralli, file)


        def Translate(input, output):
            # 转换
            wordapp = wc.Dispatch('Word.Application')
            doc = wordapp.Documents.Open(input)
            # 为了让python可以在后续操作中r方式读取txt和不产生乱码，参数为4
            doc.SaveAs(output,4 )
            doc.Close()


        input_file =  file_path
        output_file = 'D:\PyProject\docx_input\\file\\txt2\\'+file[:-5]+'.txt'
        if 0:
            Translate(input_file, output_file)

        # with open(output_file, 'r', encoding='gbk') as f:
        #     str_txt = f.read()
        # f.close()
        # refbeg_num = str_txt.rfind("参考文献")
        # print(list(str_txt[refbeg_num:].split('\n')))
        doc = Document(file_path)
        # print(file_path)
        # print(len(doc.paragraphs))
        name = file.split('.')[0]
        x=article(name,doc.paragraphs)
        # x.setname(name)
        # x.setparagraphs(doc.paragraphs)
        articlelist.append(x)
    #rws build
    for a in articlelist:
        a.setrws(rws(a.getparagraphs(),a.getname()))#任务书

        res=zykeywordextract(a.getparagraphs(),a)#摘要和关键词
        a.setzy(res[0])
        a.setkeyword(res[1])
        a.setmenu(menuextract(a.getparagraphs(),a))#目录
        a.setblocks(blockdivide(a))#按目录分块
        a.menucheck()#目录检测
        a.blockstocsv()#保存块
        a.tosents()#生成句子（如果已有，则直接读取填充）
        a.calppl()#计算ppl，如果已经计算，则直接填充
        a.cal_sents_sim()#计算句子上下句相似度，如果已经计算，则直接填充
        a.setref2()
        a.setcite()
        a.aicheck()

    return articlelist
def article_statistic_average_sentsnum_per_paragraph(article_list):
    rank_list = []
    for a in article_list:
        rank_list.append(a.average_sentsnum_per_paragraph)
        # 所有对应指标
        rank_list.sort()
    # 排序
    for a in article_list:
        inum = rank_list.index(a.average_sentsnum_per_paragraph)
        # 寻位
        a.average_sentsnum_per_paragraph_rank = (inum + 1) / len(rank_list)
def article_statistic_average_wordssnum_per_paragraph(article_list):
    rank_list = []
    for a in article_list:
        rank_list.append(a.average_wordssnum_per_paragraph)
        # 所有对应指标
        rank_list.sort()
    # 排序
    for a in article_list:
        inum = rank_list.index(a.average_wordssnum_per_paragraph)
        # 寻位
        a.average_wordssnum_per_paragraph_rank = (inum + 1) / len(rank_list)
def article_statistic_ppl(article_list):
    rank_list = []
    for a in article_list:

        rank_list.append(a.ppl)
        # 所有对应指标
        rank_list.sort()
    # 排序
    for a in article_list:
        inum = rank_list.index(a.ppl)
        # 寻位
        a.ppl_rank = (inum + 1) / len(rank_list)
        #均值
    rank_list = []

    for a in article_list:

        rank_list.append(a.ppl_var)
        # 所有对应指标
        rank_list.sort()
    # 排序
    for a in article_list:
        inum = rank_list.index(a.ppl_var)
        # 寻位
        a.ppl_var_rank = (inum + 1) / len(rank_list)
        #均值

    rank_list = []
    for a in article_list:
        rank_list.append(a.ppl_r)
        # 所有对应指标
        rank_list.sort()
    # 排序
    for a in article_list:
        inum = rank_list.index(a.ppl_r)
        # 寻位
        a.ppl_r_rank = (inum + 1) / len(rank_list)
        # 均值



def article_statistic_sents_sim(article_list):
    article_list_avg_sents_sim = []
    for a in article_list:
        article_list_avg_sents_sim.append(a.avg_sents_sim)
    #所有对应指标
    article_list_avg_sents_sim.sort()
    # 排序
    for a in article_list:
        inum=article_list_avg_sents_sim.index(a.avg_sents_sim)
        #寻位
        a.avg_sents_sim_rank=(inum+1)/len(article_list_avg_sents_sim)

    rank_list = []

    for a in article_list:
        rank_list.append(a.sents_sim_var)
        # 所有对应指标
        rank_list.sort()
    # 排序
    for a in article_list:
        inum = rank_list.index(a.sents_sim_var)
        # 寻位
        a.sents_sim_var_rank = (inum + 1) / len(rank_list)
        # 方差

    rank_list = []
    for a in article_list:
        rank_list.append(a.sents_sim_r)
        # 所有对应指标
        rank_list.sort()
    # 排序
    for a in article_list:
        inum = rank_list.index(a.sents_sim_r)
        # 寻位
        a.sents_sim_r_rank = (inum + 1) / len(rank_list)
        # 极差
def article_statistic_avg_sents_len(article_list):
    rank_list = []
    for a in article_list:
        rank_list.append(a.average_sentence_length)
# 所有对应指标
        rank_list.sort()
# 排序
    for a in article_list:
        inum = rank_list.index(a.average_sentence_length)
    # 寻位
        a.average_sentence_length_rank = (inum + 1) / len(rank_list)
def article_statistic_aicheck_res(article_list):
    for a in article_list:
        score=0
        if a.aicheck_res[0]=='ChatGPT':
            score+=-0.8*a.aicheck_res[1]
        else:
            score +=0.8 * a.aicheck_res[1]
        if a.aicheck_res[2]=='ChatGPT':
            score+=-0.5*a.aicheck_res[3]
        else:
            score +=0.5 * a.aicheck_res[3]
        a.aiscore=score
    rank_list = []
    for a in article_list:
        rank_list.append(a.aiscore)
        # 所有对应指标
        rank_list.sort()
    # 排序
    for a in article_list:
        inum = rank_list.index(a.aiscore)
        # 寻位
        a.aiscore_rank = (inum + 1) / len(rank_list)

    rank_list = []
    for a in article_list:
        if a.aicheck_res[0] == 'ChatGPT':
            a.aiscore_1 = -0.8 * a.aicheck_res[1]
        else:
            a.aiscore_1= 0.8 * a.aicheck_res[1]
        rank_list.append(a.aiscore_1)
        # 所有对应指标
    rank_list.sort()
    # 排序
    for a in article_list:
        inum = rank_list.index(a.aiscore_1)
        # 寻位
        a.aiscore_1_rank = (inum + 1) / len(rank_list)

    rank_list = []
    for a in article_list:
        if a.aicheck_res[2] == 'ChatGPT':
            a.aiscore_2 = -0.5 * a.aicheck_res[3]
        else:
            a.aiscore_2 = 0.5 * a.aicheck_res[3]
        rank_list.append(a.aiscore_2)
        # 所有对应指标
    rank_list.sort()
    # 排序
    for a in article_list:
        inum = rank_list.index(a.aiscore_2)
        # 寻位
        a.aiscore_2_rank = (inum + 1) / len(rank_list)
# def article_statistic_paragraphslennum(article_list):
#     rank_list = []
#     for a in article_list:
#         rank_list.append(a.paragraphslennum)
#         # 所有对应指标
#         rank_list.sort()
#     # 排序
#     for a in article_list:
#         inum = rank_list.index(a.paragraphslennum)
#         # 寻位
#         a.paragraphslennum_rank = (inum + 1) / len(rank_list)
# def article_statistic_paragraphslennum(article_list):
#     rank_list = []
#     for a in article_list:
#         rank_list.append(a.paragraphslens)
#         # 所有对应指标
#         rank_list.sort()
#     # 排序
#     for a in article_list:
#         inum = rank_list.index(a.paragraphslens)
#         # 寻位
#         a.paragraphslens_rank = (inum + 1) / len(rank_list)
def article_statistic_ref_sim(article_list):
    rank_list=[]
    for a in article_list:
        for i in a.ref_cite_sim:
            for j in i:
                rank_list.append(j)
        a.ref_cite_sim_avg=sum(rank_list)/len(rank_list)
        a.ref_cite_sim_var=np.var(rank_list)
        a.ref_cite_sim_r=max(rank_list)-min(rank_list)
    rank_list1 = []
    rank_list2 = []
    rank_list3 = []
    for a in article_list:
        rank_list1.append(a.ref_cite_sim_avg)
        rank_list2.append(a.ref_cite_sim_var)
        rank_list3.append(a.ref_cite_sim_r)
    rank_list1.sort()
    rank_list2.sort()
    rank_list3.sort()
    for a in article_list:
        a.ref_cite_sim_avg_rank=(rank_list1.index(a.ref_cite_sim_avg)+1)/len(article_list)
        a.ref_cite_sim_var_rank=(rank_list2.index(a.ref_cite_sim_var)+1)/len(article_list)
        a.ref_cite_sim_r_rank=(rank_list3.index(a.ref_cite_sim_r)+1)/len(article_list)
def article_statistic(article_list):
    article_statistic_ppl(article_list)
    article_statistic_sents_sim(article_list)#平均相似度rank

    article_statistic_avg_sents_len(article_list)#平均句长rank

    article_statistic_average_sentsnum_per_paragraph(article_list)
    article_statistic_average_wordssnum_per_paragraph(article_list)

    article_statistic_aicheck_res(article_list)
    article_statistic_ref_sim(article_list)
    # article_statistic_avg_paragraphs_len(article_list)
    for a in article_list:
        a.save()
bgtime=time.time()
articlelist=article_build()
print('article_build() 用时： ',time.time()-bgtime,' s')
bgtime2=time.time()
if 0:
    article_statistic(articlelist)
    print('article_statistic 用时：',time.time()-bgtime2,' s')
# articlelist[0].blockstocsv()
# articlelist[0].calppl()
# blockdivide(articlelist[3])
# articlelist[0].pplcheck()
m=1