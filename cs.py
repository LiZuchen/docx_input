
import os
import platform
from transformers import AutoTokenizer, AutoModel
import torch
from docx import Document

import re
print(torch.cuda.is_available())
print(torch.__version__)
torch.cuda.empty_cache()
# os.environ["PYTORCH_CUDA_ALLOC_CONF"] = "max_split_size_mb:128"
# MODEL_PATH = os.environ.get('MODEL_PATH', '/HOME/scw6c94/run/ac/llm/ChatGLM3/basic_demo/32k')
MODEL_PATH = os.environ.get('MODEL_PATH', '/home/liuqt/long/32k')
TOKENIZER_PATH = os.environ.get("TOKENIZER_PATH", MODEL_PATH)

tokenizer = AutoTokenizer.from_pretrained(TOKENIZER_PATH, trust_remote_code=True)
model = AutoModel.from_pretrained(MODEL_PATH, trust_remote_code=True,device_map='cuda:0').eval()
# add .quantize(bits=4, device="cuda").cuda() before .eval() to use int4 model
# must use cuda to load int4 model

os_name = platform.system()
clear_command = 'cls' if os_name == 'Windows' else 'clear'
stop_stream = False

welcome_prompt = "欢迎使用 ChatGLM3-6B 模型，输入内容即可进行对话，clear 清空对话历史，stop 终止程序"


def build_prompt(history):
    prompt = welcome_prompt
    for query, response in history:
        prompt += f"\n\n用户：{query}"
        prompt += f"\n\nChatGLM3-6B：{response}"
    return prompt

def query_zy_class(aname,history,past_key_values,stop_stream):
    f = open(
        # r"
        r'/home/liuqt/long/zy_output_files/'+aname+'.txt',
        "r", encoding='utf-8')  # 设置文件对象
    str_zy = f.read()  # 将txt文件的所有内容读入到字符串str中
    f.close()  # 将文件关闭
    # f = open(
    #     r"D:\PyProject\docx_input\prompt\log_alg.txt",
    #     "r", encoding='utf-8')  # 设置文件对象
    # str2 = f.read()  # 将txt文件的所有内容读入到字符串str中
    # f.close()

    query = "\n用户：" + str_zy + '阅读该文章摘要，说明是什么类型的论文，回答格式为：[算法类]或者[系统类]或者[理论推导类]'
    print("\nChatGLM：", end="")
    current_length = 0
    for response, history, past_key_values in model.stream_chat(tokenizer, query, history=history, top_p=1,
                                                                temperature=0.01,
                                                                past_key_values=past_key_values,
                                                                return_past_key_values=True):
        if stop_stream:
            stop_stream = False
            break
        else:
            print(response[current_length:], end="", flush=True)
            article_class=response
            current_length = len(response)
    print("")
    return article_class
def query_menu_struct_alg(aname, history, past_key_values, stop_stream):
    f = open(
        # r"
        r'/home/liuqt/long/menu_out_files/'+aname+'.txt',
        "r", encoding='utf-8')  # 设置文件对象
    str = f.read()
    mlist=eval(str)# 将txt文件的所有内容读入到字符串str中
    f.close()  # 将文件关闭
    print (mlist)
    f = open(
        # r"
        r'/home/liuqt/long/file/txt2/'+aname+'.txt',
        "r", encoding='gbk')  # 设置文件对象
    str_txt = f.read()
    # 将txt文件的所有内容读入到字符串str中
    f.close()
    # f = open(
    #     r"D:\PyProject\docx_input\prompt\log_alg.txt",
    #     "r", encoding='utf-8')  # 设置文件对象
    # str2 = f.read()  # 将txt文件的所有内容读入到字符串str中
    # f.close()
    qlist = []

    # past_key_values, history = None, []
    # global stop_stream
    # query = "\n用户：" + str + '阅读该文目录列表，说明本文所做的实验以及对应的章节，回答格式为：一系列的[xxxxx,B]其中B为所给出目录中的章节，例如[实体链接模型实验,实体链接算法研究],[关系抽取模型实验,关系抽取算法研究],[基于知识库推理的联合优化算法实验,基于知识库推理的联合优化算法设计与实现]'
    #
    # print("\nChatGLM：", end="")
    # current_length = 0
    # for response, history, past_key_values in model.stream_chat(tokenizer, query, history=history, top_p=1,
    #                                                             temperature=0.01,
    #                                                             past_key_values=past_key_values,
    #                                                             return_past_key_values=True):
    #     if stop_stream:
    #         stop_stream = False
    #         break
    #     else:
    #         print(response[current_length:], end="", flush=True)
    #
    #         current_length = len(response)
    # print("")
    # #得到实验和对应的章节
    # query = "\n用户：结合历史信息，说明本文所做的实验对应的所有章节，仅回答对应章节名称，回答格式为：[xxxxx]其中为所给出目录中的章节，例如[实体链接算法研究],[关系抽取算法研究],[基于知识库推理的联合优化算法设计与实现],[实体链接模型实验],[关系抽取模型实验],[基于知识库推理的联合优化算法实验],[实体链接任务的样本分布偏差讨论]"
    #
    #
    # print("\nChatGLM：", end="")
    # current_length = 0
    # for response, history, past_key_values in model.stream_chat(tokenizer, query, history=history, top_p=1,
    #                                                             temperature=0.01,
    #                                                             past_key_values=past_key_values,
    #                                                             return_past_key_values=True):
    #     if stop_stream:
    #         stop_stream = False
    #         break
    #     else:
    #         print(response[current_length:], end="", flush=True)
    #
    #         current_length = len(response)
    # print("")
    # #得到章节
    # rep=response
    # tlist =re.findall('\[(.*?)]',rep)
    # tlist2=[]
    # for i in tlist:
    #     if ',' in i:
    #         x=list(i.split(','))
    #         tlist2.append(x[0].replace(' ',''))
    #         tlist2.append(x[1].replace(' ',''))
    #     else:
    #         tlist2.append(i)
    # tlist2=list(set(tlist2))
    # tlist_t=[]
    # for i in tlist2:
    #     if i in str:
    #         print(i+'in menu')
    #         tlist_t.append(i)
    #     else:
    #         print(i+'not in menu')
    # #章节验证
    # tlist_t.append('致谢')
    # #添加文末
    # t_pnum_list=[]
    # for i in tlist_t:
    #     for j in mlist:
    #         if i==j[0] or i in j[0]:
    #             t_pnum_list.append([i,j[2]])
    # #生成章节段落序号
    # print(t_pnum_list)
    #所有对应实验的paragraphs序号
    docx_path='/home/liuqt/long/file/docx/'+aname+'.docx'
    docx = Document(docx_path)
    #读入原文
    # print(file_path)
    # print(len(doc.paragraphs))

    #按txt内容中出现位置划分

    # bl=blockdivide(t_pnum_list,docx.paragraphs)
    # bl2=blockdivide2(t_pnum_list,str_txt)
    # bl3=blockdivide3(t_pnum_list,str_txt)

    return [str_txt]
def query_experiment_alg(string, history, past_key_values, stop_stream, aname, save_path):
    # past_key_values, history = None, []
    # global stop_stream

    query =  string + '阅读此部分文章，说明其完成了哪些实验（不超过10个），回答格式为一系列的:[xxxx],例如[实体链接模型实验];回答实验结果是什么，给出文中的量化数据和来源。'

    print("\nChatGLM：", end="")
    current_length = 0
    for response, history, past_key_values in model.stream_chat(tokenizer, query, history=history, top_p=1,
                                                                temperature=0.01,
                                                                past_key_values=past_key_values,
                                                                return_past_key_values=True):
        if stop_stream:
            stop_stream = False
            break
        else:
            print(response[current_length:], end="", flush=True)

            current_length = len(response)
    print("")
    #内容query结束
    #指标query开始
    query =  '实验指标结果是什么，以[xxx,xxx]的格式回答。'

    print("\nChatGLM：", end="")
    current_length = 0
    for response, history, past_key_values in model.stream_chat(tokenizer, query, history=history, top_p=1,
                                                                temperature=0.01,
                                                                past_key_values=past_key_values,
                                                                return_past_key_values=True):
        if stop_stream:
            stop_stream = False
            break
        else:
            print(response[current_length:], end="", flush=True)
            with open(save_path+aname+'[指标].txt', 'a', encoding='utf-8') as f:
                print(response[current_length:],end='', file=f)

            current_length = len(response)
    print("")
    #总结query


    query = '根据以上信息对进行归纳，回答时要涵盖量化数据：针对[]问题，提出了[], 在[]上进行了[], []实验指标为[],结果表明[]。'

    print("\nChatGLM：", end="")
    current_length = 0
    for response, history, past_key_values in model.stream_chat(tokenizer, query, history=history, top_p=1,
                                                                temperature=0.01,
                                                                past_key_values=past_key_values,
                                                                return_past_key_values=True):
        if stop_stream:
            stop_stream = False
            break
        else:
            print(response[current_length:], end="", flush=True)
            with open(save_path+aname+'[归纳].txt', 'a', encoding='utf-8') as f:
                print(response[current_length:],end='', file=f)

            current_length = len(response)
    print("")
    return
def query_menu_struct_sys(aname, history, past_key_values, stop_stream):
    # f = open(r'/home/liuqt/long/menu_out_files/'+aname+'.txt',"r", encoding='utf-8')
    # # 设置文件对象
    # str = f.read()
    # mlist=eval(str)# 将txt文件的所有内容读入到字符串str中
    # f.close()  # 将文件关闭
    # #读入目录
    # print (mlist)
    #
    #
    f = open(r'/home/liuqt/long/file/txt2/'+aname+'.txt',"r", encoding='gbk')  # 设置文件对象
    str_txt = f.read()
    # 将txt文件的所有内容读入到字符串str中
    f.close()
    # #读入文章txt版本
    #
    # # past_key_values, history = None, []
    # # global stop_stream
    # #query
    # query = "\n用户：" + str + '阅读该文目录列表，说明本文所做系统的模块以及对应的章节，回答格式为：一系列的[xxxxx,B]其中B为所给出目录中的章节，例如[xxx,xxxxx],[xxx,xxxxx]'
    # print(query)
    # print("\nChatGLM：", end="")
    # current_length = 0
    # for response, history, past_key_values in model.stream_chat(tokenizer, query, history=history, top_p=1,
    #                                                             temperature=0.01,
    #                                                             past_key_values=past_key_values,
    #                                                             return_past_key_values=True):
    #     if stop_stream:
    #         stop_stream = False
    #         break
    #     else:
    #         print(response[current_length:], end="", flush=True)
    #
    #         current_length = len(response)
    # print("")
    # #得到实验和对应的章节
    # query = "\n用户：结合历史信息，说明本文涉及到的章节名称，只回答对应章节名称，回答格式例如[系统性能需求分析],[质量模块需求分析]，一个中括号中只能有一个章节名称"
    # print(query)
    #
    # print("\nChatGLM：", end="")
    # current_length = 0
    # for response, history, past_key_values in model.stream_chat(tokenizer, query, history=history, top_p=1,
    #                                                             temperature=0.01,
    #                                                             past_key_values=past_key_values,
    #                                                             return_past_key_values=True):
    #     if stop_stream:
    #         stop_stream = False
    #         break
    #     else:
    #         print(response[current_length:], end="", flush=True)
    #
    #         current_length = len(response)
    # print("")
    # #得到章节
    # rep=response
    # tlist =re.findall('\[(.*?)]',rep)
    # tlist_t=[]
    # for i in tlist:
    #     if i in str:
    #         print(i+'in menu')
    #         tlist_t.append(i)
    #     else:
    #         print(i+'not in menu')
    # #章节验证
    # tlist_t.append('致谢')
    # print('tlist_t')
    # print(tlist_t)
    # print('mlist')
    # print(mlist)
    # #添加文末
    # t_pnum_list=[]
    # for i in tlist_t:
    #     for j in mlist:
    #         #可能是i==j[0]
    #         if i in j[0]:
    #             t_pnum_list.append([i,j[2]])
    # #生成章节段落序号
    # print(t_pnum_list)
    #所有对应实验的paragraphs序号
    docx_path='/home/liuqt/long/file/docx/'+aname+'.docx'
    docx = Document(docx_path)
    #读入原文
    # print(file_path)
    # print(len(doc.paragraphs))

    #按txt内容中出现位置划分
    def blockdivide2(hd,str_txt):
        # print(a.getname())
        blocks = []
        hdnum = []
        # menuitem===>text ,heading x, pi
        for i in hd:
            hdnum.append(str_txt.find(i[0]))  # hd[2]--->pi
        print(hdnum.sort())
        for num in range(0, len(hdnum) - 1):
            blocklist = []
            if num < len(hdnum) - 1:
                # for i in str_txt[hdnum[num]:hdnum[num + 1]]:
                #     # print(i.text)
                blocklist.append(str_txt[hdnum[num]:hdnum[num + 1]])
            # else:
            #     for i in a.getparagraphs()[hdnum[num]:len(a.getparagraphs())]:
            #         # print(i.text)
            #         blocklist.append(i.text)
            # print("----------------------------------------------------------------------")
            blocks.append([hdnum[num], hdnum[num + 1], blocklist])
        return blocks

    #按docx内段落划分
    # def blockdivide(hd,p):
    #     # print(a.getname())
    #     blocks = []
    #     hdnum = []
    #     # menuitem===>text ,heading x, pi
    #     for i in hd:
    #         hdnum.append(i[1])  # hd[2]--->pi
    #     hdnum.sort()
    #     for num in range(0, len(hdnum)-1):
    #         blocklist = []
    #         if num < len(hdnum) - 1:
    #             for i in p[hdnum[num]:hdnum[num + 1]]:
    #                 # print(i.text)
    #                 blocklist.append(i.text)
    #         # else:
    #         #     for i in a.getparagraphs()[hdnum[num]:len(a.getparagraphs())]:
    #         #         # print(i.text)
    #         #         blocklist.append(i.text)
    #         # print("----------------------------------------------------------------------")
    #         blocks.append([hdnum[num], hdnum[num + 1], blocklist])
    #     return blocks
    # bl=blockdivide(t_pnum_list,docx.paragraphs)
    # bl2=blockdivide2(t_pnum_list,str_txt)
    # bl3=blockdivide3(t_pnum_list,str_txt)

    return [str_txt]
def query_experiment_sys(string, history, past_key_values, stop_stream, aname, save_path):
    # past_key_values, history = None, []
    # global stop_stream

    query =  string + '阅读此部分文章，回答本文所实现的系统或软件有哪些模块（不超过10个），回答格式为[xxx],[xxx],[xxx]，例如[用户管理模块],[题目管理模块]。'

    print("\nChatGLM：", end="")
    current_length = 0
    for response, history, past_key_values in model.stream_chat(tokenizer, query, history=history, top_p=1,
                                                                temperature=0.01,
                                                                past_key_values=past_key_values,
                                                                return_past_key_values=True):
        if stop_stream:
            stop_stream = False
            break
        else:
            print(response[current_length:], end="", flush=True)
            rep=response
            current_length = len(response)
    print("")
    #内容query结束
    #指标query开始

    query = '回答每一个模块的需求是什么，架构是什么，功能是什么，该模块测试方式是什么,如果有对应指标，给出所在文中的依据，否则回答[未找到具体指标]，回答格式为：[xxx]。'
    print(query)
    print("\nChatGLM：", end="")
    current_length = 0
    for response, history, past_key_values in model.stream_chat(tokenizer, query, history=history, top_p=1,
                                                                    temperature=0.01,
                                                                    past_key_values=past_key_values,
                                                                    return_past_key_values=True):
        if stop_stream:
            stop_stream = False
            break
        else:
            print(response[current_length:], end="", flush=True)
            with open(save_path+aname+'[指标].txt', 'a', encoding='utf-8') as f:
                print(response[current_length:],end='', file=f)
            f.close()
            current_length = len(response)
    print("")

    #总结query
    return
def query_summary_sys(aname, history, past_key_values, stop_stream,save_path):
    query = '根据以上信息对每一个模块进行归纳，回答时尽可能涵盖量化数据，回答格式为：针对[]需求，提出了[]模块, 进行了[], 结果表明[]。'
    print("\nChatGLM：", end="")
    current_length = 0
    for response, history, past_key_values in model.stream_chat(tokenizer, query, history=history, top_p=1,
                                                                temperature=0.01,
                                                                past_key_values=past_key_values,
                                                                return_past_key_values=True):
        if stop_stream:
            stop_stream = False
            break
        else:
            print(response[current_length:], end="", flush=True)
            with open(save_path + aname + '[归纳].txt', 'a', encoding='utf-8') as f:
                print(response[current_length:], end='', file=f)
            f.close()
            current_length = len(response)
    print("")
def blockdivide2(hd,str_txt):
        # print(a.getname())
    blocks = []
    hdnum = []
    # menuitem===>text ,heading x, pi
    for i in hd:
        print(i,str_txt.find(i[0]))
        hdnum.append(str_txt.find(i[0]))
            # hd[2]--->pi
    print(hdnum.sort())
    for num in range(0, len(hdnum) - 1):
        blocklist = []
        if num < len(hdnum) - 1:
                # for i in str_txt[hdnum[num]:hdnum[num + 1]]:
                #     # print(i.text)
            blocklist.append(str_txt[hdnum[num]:hdnum[num + 1]])
            # else:
            #     for i in a.getparagraphs()[hdnum[num]:len(a.getparagraphs())]:
            #         # print(i.text)
            #         blocklist.append(i.text)
            # print("----------------------------------------------------------------------")
        blocks.append([hdnum[num], hdnum[num + 1], blocklist])
    return blocks

def blockdivide3(hd, str_txt):
        # return [str_txt[0:int(len(str_txt)/2)],str_txt[int(len(str_txt)/2):]]
    return [str_txt]
    #按docx内段落划分
def blockdivide(hd,p):
        # print(a.getname())
        blocks = []
        hdnum = []
        # menuitem===>text ,heading x, pi
        for i in hd:
            hdnum.append(i[1])  # hd[2]--->pi
        hdnum.sort()
        for num in range(0, len(hdnum)-1):
            blocklist = []
            if num < len(hdnum) - 1:
                for i in p[hdnum[num]:hdnum[num + 1]]:
                    # print(i.text)
                    blocklist.append(i.text)
            # else:
            #     for i in a.getparagraphs()[hdnum[num]:len(a.getparagraphs())]:
            #         # print(i.text)
            #         blocklist.append(i.text)
            # print("----------------------------------------------------------------------")
            blocks.append([hdnum[num], hdnum[num + 1], blocklist])
        return blocks
def main(args):
    for i in args[0]:
        past_key_values, history = None, []
        global stop_stream
        save_path2='/home/liuqt/long/file/logsave/'
        save_path = 'D:\PyProject\docx_input\\file\logsave\\'
        # aname='17373118_刘阳_面向国产ARM平台的浮点计算分析工具（最终版）'
        aname=i
        # aname='17373086_黎明_C语言程序自动评测系统设计与实现'
        articleclass=query_zy_class(aname,history,past_key_values,stop_stream)
        if '算法类' in articleclass:
            with open(save_path2 + aname + '[归纳].txt', 'w', encoding='utf-8') as f:
                print('文章归纳', end='\n', file=f)
            f.close()
            with open(save_path2 + aname + '[指标].txt', 'w', encoding='utf-8') as f:
                print('文章指标', end='\n', file=f)
            f.close()
            bl=query_menu_struct_alg(aname, history, past_key_values, stop_stream)
            query_experiment_alg(bl[0], history, past_key_values, stop_stream, aname, save_path2)

            # for i in bl:
            #     print(i[2])
            #     string=';'.join(i[2])
            #     query_experiment_alg(string, history, past_key_values, stop_stream, aname, save_path2)
        elif '系统类' in articleclass:
            bl = query_menu_struct_sys(aname, history, past_key_values, stop_stream)
            # for i in bl:
            #     print(i[2])
            #     string = ';'.join(i[2])
            #clear
            with open(save_path2 + aname + '[归纳].txt', 'w', encoding='utf-8') as f:
                print('文章归纳', end='\n', file=f)
            f.close()
            with open(save_path2 + aname + '[指标].txt', 'w', encoding='utf-8') as f:
                print('文章指标', end='\n', file=f)
            f.close()
            query_experiment_sys(bl[0], history, past_key_values, stop_stream, aname, save_path2)
            query_summary_sys(aname,history,past_key_values,stop_stream,save_path2)

        else:
            print('理论推导类')


    # while True:
    #     query = input("\n用户：")
    #     if query.strip() == "stop":
    #         break
    #     if query.strip() == "clear":
    #         past_key_values, history = None, []
    #         os.system(clear_command)
    #         print(welcome_prompt)
    #         continue
    #     print("\nChatGLM：", end="")
    #     current_length = 0
    #     for response, history, past_key_values in model.stream_chat(tokenizer, query, history=history, top_p=1,
    #                                                                 temperature=0.01,
    #                                                                 past_key_values=past_key_values,
    #                                                                 return_past_key_values=True):
    #         if stop_stream:
    #             stop_stream = False
    #             break
    #         else:
    #             print(response[current_length:], end="", flush=True)
    #             current_length = len(response)
    #     print("")


if __name__ == "__main__":
    files=os.listdir('file/txt2')
    filenames=[]
    for i in files:
        filenames.append(i.split('.')[0])
    args=[filenames]

    main(args)
